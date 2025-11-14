import os
import io
import json
import subprocess
import logging
import re
import sys
import unicodedata
from datetime import datetime, timedelta
from dateutil import parser as date_parser
from docx import Document
from werkzeug.utils import secure_filename
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2.credentials import Credentials
from models.models_usuario import get_db_session, Usuario
import whisper

# ----------------------------
# Logger
# ----------------------------
os.makedirs("logs", exist_ok=True)
logging.basicConfig(
    filename="logs/atas.log",
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%d/%m/%Y %H:%M:%S"
)

# ----------------------------
# Whisper - inicialização segura
# ----------------------------
try:
    # Detecta se está rodando empacotado com PyInstaller
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")

    modelo_path = os.path.join(base_path, "small.pt")

    logging.info(f"Carregando modelo Whisper de: {modelo_path}")

    if os.path.exists(modelo_path):
        whisper_model = whisper.load_model(modelo_path)
    else:
        whisper_model = whisper.load_model("small")

    logging.info("Modelo Whisper carregado com sucesso.")
except Exception as e:
    logging.error(f"Falha ao carregar modelo Whisper: {e}")
    whisper_model = None

# ----------------------------
# Status de processamento de ATA
# ----------------------------
ata_status = {}

def log_status(event_id, mensagem, erro=False):
    ata_status[event_id]["mensagem"] = mensagem
    if erro:
        ata_status[event_id]["erro"] = True
        logging.error(f"[{event_id}] {mensagem}")
    else:
        logging.info(f"[{event_id}] {mensagem}")

# ----------------------------
# Usuário
# ----------------------------
def login_user(email, senha):
    session = get_db_session()
    try:
        user = session.query(Usuario).filter_by(email=email, senha=senha).first()
        if user:
            session.expunge(user)
            return user
        return None
    finally:
        session.close()


def cadastrar_usuario(nome, email, senha, arquivo=None):
    session = get_db_session()
    try:
        if session.query(Usuario).filter_by(email=email).first():
            return None

        # Cria o usuário salvando a senha em texto puro
        user = Usuario(nome=nome, email=email, senha=senha)
        session.add(user)
        session.commit()

        # Caso haja arquivo, salva também
        if arquivo:
            user_folder = get_user_upload_folder(email)
            destino = os.path.join(user_folder, secure_filename(os.path.basename(arquivo)))
            import shutil
            shutil.copy(arquivo, destino)

        session.expunge(user)
        return user
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()

# ----------------------------
# Pastas do usuário
# ----------------------------
def get_user_upload_folder(email):
    if callable(email):
        email = email()
    email = str(email)
    safe_email = email.replace("@", "_").replace(".", "_")

    folder = os.path.join("uploads", safe_email)

    # Cria apenas se não existir (evita erro em sistemas de rede)
    if not os.path.exists(folder):
        try:
            os.makedirs(folder, exist_ok=True)
        except FileExistsError:
            pass  # outro processo pode ter criado ao mesmo tempo
        except PermissionError:
            logging.warning(f"Sem permissão para criar a pasta {folder}")

    # Define permissões apenas se for local (não rede)
    if os.path.isdir(folder) and not folder.startswith("\\\\"):
        try:
            os.chmod(folder, 0o777)
        except Exception as e:
            logging.warning(f"Não foi possível alterar permissões de {folder}: {e}")

    return folder


def get_user_download_folder(email):
    safe_email = email.replace("@", "_").replace(".", "_")
    folder = os.path.join("downloads", safe_email)
    os.makedirs(folder, exist_ok=True)
    os.chmod(folder, 0o777)
    return folder

# ----------------------------
# Credenciais Google
# ----------------------------
def salvar_credenciais_usuario(user_id, credenciais):
    session = get_db_session()
    try:
        user = session.query(Usuario).filter_by(id=user_id).first()
        if not user:
            raise ValueError("Usuário não encontrado.")

        cred_dict = {
            'token': credenciais.token,
            'refresh_token': credenciais.refresh_token,
            'token_uri': credenciais.token_uri,
            'client_id': credenciais.client_id,
            'client_secret': credenciais.client_secret,
            'scopes': credenciais.scopes
        }
        user.google_token = json.dumps(cred_dict)
        session.commit()
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()

def carregar_credenciais_usuario(user_id):
    session = get_db_session()
    try:
        user = session.query(Usuario).filter_by(id=user_id).first()
        if not user or not user.google_token:
            return None

        cred_dict = json.loads(user.google_token)
        creds = Credentials(
            token=cred_dict.get('token'),
            refresh_token=cred_dict.get('refresh_token'),
            token_uri=cred_dict.get('token_uri'),
            client_id=cred_dict.get('client_id'),
            client_secret=cred_dict.get('client_secret'),
            scopes=cred_dict.get('scopes')
        )
        return creds
    finally:
        session.close()

# ----------------------------
# Google Calendar / Drive
# ----------------------------
def listar_reunioes(creds):
    service = build('calendar', 'v3', credentials=creds)
    all_events = []
    page_token = None
    while True:
        events_result = service.events().list(
            calendarId='primary',
            maxResults=2500,
            singleEvents=True,
            orderBy='startTime',
            pageToken=page_token
        ).execute()
        events = events_result.get('items', [])
        all_events.extend(events)
        page_token = events_result.get('nextPageToken')
        if not page_token:
            break

    meet_events = [
        e for e in all_events
        if 'hangoutLink' in e or (e.get('conferenceData') and e['conferenceData'].get('entryPoints'))
    ]
    return meet_events

def baixar_video_drive(creds, file_id, destino):
    service = build('drive', 'v3', credentials=creds)
    request = service.files().get_media(fileId=file_id)
    with io.FileIO(destino, 'wb') as fh:
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()

# ----------------------------
# Limpeza de nomes de arquivos
# ----------------------------
def limpar_nome_arquivo(nome):
    nome_sem_acentos = ''.join(
        c for c in unicodedata.normalize('NFD', nome)
        if unicodedata.category(c) != 'Mn'
    )
    nome_limpo = re.sub(r'[^a-zA-Z0-9._-]', '_', nome_sem_acentos)
    return nome_limpo

# ----------------------------
# Extração de áudio com FFmpeg
# ----------------------------
def extrair_audio_ffmpeg(video_path, audio_path):
    try:
        # Garante que o vídeo existe
        if not os.path.exists(video_path):
            raise FileNotFoundError(f"Vídeo não encontrado: {video_path}")

        # Limpa apenas os nomes dos arquivos
        video_dir = os.path.dirname(video_path)
        audio_dir = os.path.dirname(audio_path)
        base_name = os.path.splitext(os.path.basename(video_path))[0]

        # Gera nomes limpos e seguros
        base_name_limpo = limpar_nome_arquivo(base_name)
        video_path_limpo = os.path.join(video_dir, f"{base_name_limpo}.mp4")
        audio_path_limpo = os.path.join(audio_dir, f"{base_name_limpo}.wav")

        # Renomeia o vídeo apenas se necessário
        if video_path != video_path_limpo:
            try:
                os.rename(video_path, video_path_limpo)
                video_path = video_path_limpo
            except FileExistsError:
                video_path = video_path_limpo  # usa o existente
            except Exception as e:
                print(f"⚠️ Aviso: não foi possível renomear o vídeo ({e}). Continuando...")

        # Executa o ffmpeg para extrair o áudio
        subprocess.run([
            "ffmpeg", "-y", "-i", video_path,
            "-vn", "-acodec", "pcm_s16le",
            "-ar", "16000", "-ac", "1",
            "-f", "wav", audio_path_limpo
        ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        # Confirma que o arquivo foi criado
        if not os.path.exists(audio_path_limpo):
            raise FileNotFoundError(f"Falha ao criar o arquivo de áudio: {audio_path_limpo}")

        return audio_path_limpo

    except subprocess.CalledProcessError as e:
        print(f"❌ Erro ao extrair áudio com ffmpeg: {e}")
        raise
    except Exception as e:
        print(f"❌ Erro inesperado ao extrair áudio: {e}")
        raise


# ----------------------------
# Buscar senha (exemplo)
# ----------------------------
def buscar_senha_por_email(email):
    session = get_db_session()
    usuario = session.query(Usuario).filter_by(email=email).first()
    session.close()
    if usuario:
        return usuario.senha
    return None

# ----------------------------
# Geração de resumo
# ----------------------------
def gerar_resumo_texto(texto):
    if not texto:
        return "Resumo indisponível. A ata não contém texto."

    texto = re.sub(r'\s+', ' ', texto.strip())
    frases = re.split(r'(?<=[.!?]) +', texto)

    if len(frases) <= 5:
        return texto

    palavras_chave = ['decisão', 'conclusão', 'ação', 'responsável', 'pauta', 'acordo']
    frases_relevantes = [f for f in frases if any(p in f.lower() for p in palavras_chave)]

    if len(frases_relevantes) < 3:
        frases_relevantes = frases[:5]

    resumo = ' '.join(frases_relevantes[:6])
    return resumo.strip()

# ----------------------------
# Processamento de ATA com Whisper
# ----------------------------
def processar_ata(event, creds, usuario_email, callback=None):
    event_id = event.get('id')
    ata_status[event_id] = {"ready": False, "erro": False, "mensagem": "Iniciando..."}

    try:
        usuario_email = usuario_email() if callable(usuario_email) else str(usuario_email)
        user_folder = get_user_upload_folder(usuario_email)
        log_status(event_id, "Baixando vídeo do Google Drive...")

        transcricao = ""

        if creds:
            service = build('drive', 'v3', credentials=creds)
            data_evento = date_parser.parse(event['start']['dateTime']).date()

            # Busca vídeos ±1 dia (para compensar fuso horário)
            inicio = (data_evento - timedelta(days=1)).strftime("%Y-%m-%dT00:00:00")
            fim = (data_evento + timedelta(days=1)).strftime("%Y-%m-%dT23:59:59")

            query = (
                f"mimeType='video/mp4' and createdTime >= '{inicio}' "
                f"and createdTime <= '{fim}'"
            )

            results = service.files().list(
                q=query,
                pageSize=1,
                orderBy='createdTime desc',
                fields="files(id, name)"
            ).execute()
            files = results.get('files', [])

            if files:
                video_file = files[0]

                # --- Limpeza e padronização do nome ---
                nome_original = os.path.splitext(video_file['name'])[0]
                nome_limpo = limpar_nome_arquivo(nome_original)
                extensao_video = os.path.splitext(video_file['name'])[1] or ".mp4"

                nome_video = os.path.join(user_folder, f"{nome_limpo}{extensao_video}")
                audio_file = os.path.join(user_folder, f"{nome_limpo}.wav")

                # --- Evita sobrescritas automáticas ---
                contador = 1
                while os.path.exists(nome_video) or os.path.exists(audio_file):
                    nome_video = os.path.join(user_folder, f"{nome_limpo}_{contador}{extensao_video}")
                    audio_file = os.path.join(user_folder, f"{nome_limpo}_{contador}.wav")
                    contador += 1

                # --- Logs ---
                log_status(event_id, f"Baixando vídeo do Drive: {video_file['name']}")
                log_status(event_id, f"Salvando como: {os.path.basename(nome_video)}")

                # --- Download seguro do vídeo ---
                baixar_video_drive(creds, video_file['id'], nome_video)

                # --- Extração de áudio ---
                log_status(event_id, f"Extraindo áudio de: {nome_video} → {audio_file}")
                extrair_audio_ffmpeg(nome_video, audio_file)

                # (Opcional) Remove o .mp4 após extração, para economizar espaço
                try:
                    os.remove(nome_video)
                    log_status(event_id, "Vídeo removido após extração para liberar espaço.")
                except Exception as e_rm:
                    log_status(event_id, f"Falha ao remover vídeo: {e_rm}")

                # --- Transcrição com Whisper ---
                if whisper_model is None:
                    raise RuntimeError("O modelo Whisper não foi carregado corretamente.")

                log_status(event_id, "Transcrevendo áudio...")
                result = whisper_model.transcribe(audio_file, language='pt')
                transcricao = result.get('text', '')

            else:
                transcricao = "Nenhum vídeo encontrado no Google Drive para este evento."
        else:
            transcricao = "Transcrição simulada (sem credenciais Google)."

        # --- Geração do resumo e DOCX ---
        resumo = gerar_resumo_texto(transcricao)
        ata_status[event_id]["resumo"] = resumo

        log_status(event_id, "Gerando DOCX...")
        doc = Document()
        doc.add_heading('ATA DE REUNIÃO', 0)
        doc.add_paragraph(f"Data: {datetime.today().strftime('%d/%m/%Y')}")
        doc.add_paragraph(f"Reunião: {event.get('summary', '')}")
        doc.add_heading('Resumo', level=1)
        doc.add_paragraph(resumo)
        doc.add_heading('Transcrição', level=1)
        doc.add_paragraph(transcricao)

        caminho = os.path.join(
            user_folder,
            f'ata_{event_id}_{datetime.now().strftime("%H%M%S")}.docx'
        )
        doc.save(caminho)

        ata_status[event_id].update({"ready": True, "mensagem": f"Concluído! Arquivo: {caminho}"})

    except Exception as e:
        log_status(event_id, f"Erro ao processar ATA: {e}", erro=True)
    finally:
        if callback:
            caminho_final = caminho if ata_status[event_id]["ready"] else None
            callback(event_id, caminho_final)

